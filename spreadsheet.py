from __future__ import print_function
import pickle
import os.path
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow, Flow
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from google.auth.transport.requests import Request
from pprint import pprint
from googleapiclient import discovery
import docx
from docx.shared import Pt
import base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import mimetypes

# If modifying these scopes, delete the file token.pickle
SCOPES = ['https://www.googleapis.com/auth/drive']

# The ID and range of a sample spreadsheet
dictOfSpreadsheetsId = {"Group1":"/insert-spreadsheet-id", # insert!
"Group2":"/insert-spreadsheet-id",
"Group3":"/insert-spreadsheet-id"
}
SAMPLE_SPREADSHEET_ID = ''
#SAMPLE_RANGE_NAME = 'sheet1!B2:AV60'
request_body = {
    'requests': [
    {
        'duplicateSheet':{
            'sourceSheetId': '',
            'newSheetName': ''
        }
    }]
}
def main():
    creds = None
    # The file token.pickle stores the user's access and refresh tokens, and is created automatically when the authorization flow completes for the first time
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid: # If there are no (valid) credentials available, let the user log in
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.pickle', 'wb') as token:  # Save the credentials for the next run
            pickle.dump(creds, token)
    service = discovery.build('sheets', 'v4', credentials=creds)

# Create service to connect to spreadsheet by using Google API
    def Create_Service(client_secret_file, api_name, api_version, *scopes):
        CLIENT_SECRET_FILE = client_secret_file
        API_SERVICE_NAME = api_name
        API_VERSION = api_version
        SCOPES = [scope for scope in scopes[0]]

        cred = None

        pickle_file = f'token_{API_SERVICE_NAME}_{API_VERSION}.pickle'
        # print(pickle_file)

        if os.path.exists(pickle_file):
            with open(pickle_file, 'rb') as token:
                cred = pickle.load(token)

        if not cred or not cred.valid:
            if cred and cred.expired and cred.refresh_token:
                cred.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(CLIENT_SECRET_FILE, SCOPES)
                cred = flow.run_local_server()

            with open(pickle_file, 'wb') as token:
                pickle.dump(cred, token)

        try:
            service = build(API_SERVICE_NAME, API_VERSION, credentials=cred)
            return service
        except Exception as e:
            print('Unable to connect.')
            print(e)
            return None   

# Send e-mail with file that has order information from spreadsheet
    def sendMail():
        CLIENT_SECRET_FILE = 'client_secret.json'
        API_NAME = 'gmail'
        API_VERSION = 'v1'
        SCOPES = ['https://mail.google.com/']

        service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)
        service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)
         
        file_attachments = [r'my_doc_file.docx'] # File You want to send
         
        emailMsg = 'One file attached'
         
        # Create email message
        mimeMessage = MIMEMultipart()
        mimeMessage['to'] = 'insert!' # Put here the receiver e-mail
        mimeMessage['subject'] = 'You got files'
        mimeMessage.attach(MIMEText(emailMsg, 'plain'))
         
        # Attach files
        for attachment in file_attachments:
            content_type, encoding = mimetypes.guess_type(attachment)
            main_type, sub_type = content_type.split('/', 1)
            file_name = os.path.basename(attachment)
         
            f = open(attachment, 'rb')
         
            myFile = MIMEBase(main_type, sub_type)
            myFile.set_payload(f.read())
            myFile.add_header('Content-Disposition', 'attachment', filename=file_name)
            encoders.encode_base64(myFile) 
            f.close()   
            mimeMessage.attach(myFile)
         
        raw_string = base64.urlsafe_b64encode(mimeMessage.as_bytes()).decode()
         
        message = service.users().messages().send(
            userId='me',
            body={'raw': raw_string}).execute()

# Get data from spreadsheet with order info
    def getDataFromSheet():
        sheet = service.spreadsheets() # Call the Sheets API
        global newSheetName
        newSheetName = request_body['requests'][0]['duplicateSheet']['newSheetName']
        if newSheetName == '':
            print('Insert the name of sheet to get data from:')
            newSheetName = str(input())
        SAMPLE_RANGE_NAME = ('%s!B2:AV60' %newSheetName)
        result = sheet.values().get(spreadsheetId=SAMPLE_SPREADSHEET_ID, range=SAMPLE_RANGE_NAME).execute()
        values = result.get('values', [])
        namesOfProducts = values[1][5::]
        prizesOfProductsInZloty = values[4][5::]
        prizesOfProducts = [] # List of turned string prizes to int
        for prizeOfProductInZloty in prizesOfProductsInZloty:
            prizeOfProduct = int(''.join(list(prizeOfProductInZloty)[0:-6])) # Delete ',00 zł'' then turn from list to int
            prizesOfProducts.append(prizeOfProduct)
        productsDescription1 = values[2][5::]
        productsDescription2 = values[3][5::]
        quantityOfProducts = values[5][5::]
        dictOfBoughtProducts = {}
        for i in range(0,len(quantityOfProducts)):
            if (quantityOfProducts[i] !='0' and quantityOfProducts[i] !='0,0'):
                dictOfBoughtProducts[namesOfProducts[i]] = quantityOfProducts[i]
        totalAmmountAllOrders = values[0][4]
        ammountForPurchaser = values[4][5::]
        global dictOfOrderInfoAllBuyers    
        dictOfOrderInfoAllBuyers = {}
        dictOfProductAndQuantity = {}
        for i in range(0, 40):
            listOfAllProductsForBuyer = []
            nameOfBuyer = values[i+6][0]
            if nameOfBuyer == '':
                continue
            for j in range(0, len(quantityOfProducts)):
                nameOfProduct = values[1][j+5]
                quantityOfProduct = values[i+6][j+5]
                if quantityOfProduct == '':
                    continue
                if nameOfProduct == '99,00 zł': # At the final column must something so as to create the same lenghth of elements in values
                    break                    
                dictOfProductAndQuantity[nameOfProduct] = quantityOfProduct
                listOfAllProductsForBuyer.append(dictOfProductAndQuantity)
                dictOfProductAndQuantity = {}
                dictOfOrderInfoAllBuyers[nameOfBuyer] = listOfAllProductsForBuyer # Dict with lists of dict {'nick1':[{'Product1':'Quantity1'},...], 'nick2' ....} 
        counter = 0
        print('Print orders for buyers? y/n')
        ifDisplayOrderInfo = str(input())
        if ifDisplayOrderInfo == 'y':
            for i in dictOfOrderInfoAllBuyers:
                counter = counter +1
                print('\nnr.%s.' %counter)
                print('%s'%i)
                list1 = dictOfOrderInfoAllBuyers[i]
                for j in list1:
                    print(j)

# Create .docx file with order info
    def createDocx():
        mydoc = docx.Document()
        style = mydoc.styles['Normal']
        buyerCounter = 0
        paragraph = mydoc.add_paragraph(str(newSheetName))
        for nick in dictOfOrderInfoAllBuyers:
            buyerCounter = buyerCounter +1
            paragraph = mydoc.add_paragraph(str('nr.%s.'%buyerCounter))
            paragraph = mydoc.add_paragraph(str(nick))
            for productName in range (0, len(dictOfOrderInfoAllBuyers[nick])):
                for i in dictOfOrderInfoAllBuyers[nick][productName]:
                    productInfo = ('%s x %s' % (i, (dictOfOrderInfoAllBuyers[nick][productName][i]) ))
                    paragraph = mydoc.add_paragraph(str(productInfo))
            paragraph = mydoc.add_paragraph('')
        font = style.font
        font.name = 'Arial'
        font_size = Pt(5)
        paragraph.style = mydoc.styles['Normal']
        mydoc.save('my_doc_file.docx')
        print('The .docx file was created.')

# Duplicate sheet to create a new sheet
    def duplicateSheet():
        request_body['requests'][0]['duplicateSheet']['newSheetName'] = ('Smakołyki %s.%s.%s %s' %(eventDay, eventMonth, eventYear, eventGroup))
        request_body['requests'][0]['duplicateSheet']['sourceSheetId'] = (sheetIdToCopy) # get sheetIdToCopy from createEvent()
        #newSheetName = request_body['requests'][0]['duplicateSheet']['newSheetName']
        response = service.spreadsheets().batchUpdate(
            spreadsheetId = SAMPLE_SPREADSHEET_ID,
            body = request_body
        ).execute()

# Color column so as to mark orders with suceeded payment, or overdue with payment
    def colorColumn(columnToColor, newSheetId):
        request_body_color={
        'requests':[
        {
        'repeatCell':{
            'range':{
                'sheetId': newSheetId,
                'startRowIndex': columnToColor,
                'endRowIndex': columnToColor+1,
                'startColumnIndex':0,
            },
        'cell':{
            "userEnteredFormat": {
            "backgroundColor": {
              "green": 0,
              'red': 0
            }
          }
        },
        "fields": "userEnteredFormat(backgroundColor)"
      }}]}
        print('Choose color to color row r/y/g:')
        redOrGreen = str(input())
        if redOrGreen == 'r':
            request_body_color['requests'][0]['repeatCell']['cell']['userEnteredFormat']['backgroundColor']['red'] = 1
        if redOrGreen == 'g':
            request_body_color['requests'][0]['repeatCell']['cell']['userEnteredFormat']['backgroundColor']['green'] = 1
        if redOrGreen == 'y':
            request_body_color['requests'][0]['repeatCell']['cell']['userEnteredFormat']['backgroundColor']['red'] = 1
            request_body_color['requests'][0]['repeatCell']['cell']['userEnteredFormat']['backgroundColor']['green'] = 1
        response = service.spreadsheets().batchUpdate(
            spreadsheetId = SAMPLE_SPREADSHEET_ID,
            body = request_body_color
        ).execute()

# Clear values of range of cells in newly created sheet, that was copied from chosen sheet
    def clearValuesOfSpreadsheet():
        spreadsheetIdToClear = SAMPLE_SPREADSHEET_ID
        clear_values_request_body = {
        }
        def requestClearValues(rangeToClear):
            request = service.spreadsheets().values().clear(spreadsheetId=spreadsheetIdToClear,range=rangeToClear, body=clear_values_request_body)
            response = request.execute()        
        # Loop so as to choose new range of cells to clear
        while True:
            print('\nInsert range to clear: (for example: B8:E40 ; G8:AT40)')
            rangeToClear = str(input())
            requestClearValues(rangeToClear)
            print('Clear more? y/n')
            ifAgainSetRangeToClear = str(input())
            if ifAgainSetRangeToClear == 'y':
                continue
            else:
                break
        print('Specified range was cleaned.')

# Choose group to operate on it
    def chooseGroup():
        global eventGroup
        print('Enter group to operate on: Group1/Group2/Group3')
        eventGroup = str(input())
        global SAMPLE_SPREADSHEET_ID
        if eventGroup == 'Group1':
            SAMPLE_SPREADSHEET_ID = dictOfSpreadsheetsId["Group1"]
        if eventGroup == 'Group2':
            SAMPLE_SPREADSHEET_ID = dictOfSpreadsheetsId["Group2"]
        if eventGroup == 'Group3':
            SAMPLE_SPREADSHEET_ID = dictOfSpreadsheetsId["Group3"]

# Create event description
    def createEvent():
        global sheetIdToCopy
        print('Enter spreadsheet ID to copy:')
        sheetIdToCopy = str(input())
        print('Enter the day of event: (w-wednesday, t-thursday)')
        eventDayCheck = str(input())
        print(' Enter the day of event: (for example 2)')
        eventDay = int(input())
        print('Enter the month of event:')
        eventMonth = int(input())
        print('Enter the year of event:')
        eventYear = int(input())

        # Check if leap year and check numbers of days so as to define order and payment day
        month_days = [0, 31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
        def is_leap(year):
            return year % 4 == 0 and (year % 100 != 0 or year % 400 == 0)

        def days_in_month(year, month):
            if not 1 <= month <= 12:
                return 'Wrong month.'

            if month == 2 and is_leap(year):
                return 29
            return month_days[month]

        def event_order_payment(eventDayCheck, eventDay, eventMonth):
            if eventDayCheck == "w":
                orderDay = eventDay - 4
            elif eventDayCheck == "t":
                orderDay = eventDay - 5

            def month_before_event(eventMonth, orderDay):   
                if eventMonth < 1:
                    monthBeforeEvent = 12
                else:
                    monthBeforeEvent = eventMonth - 1
                daysInMonthBeforeEvent = days_in_month(eventYear, monthBeforeEvent)
                orderDay = daysInMonthBeforeEvent + orderDay
                return orderDay, monthBeforeEvent
            if orderDay <=0:
                orderDay = month_before_event(eventMonth, orderDay)[0]
                orderAndPaymentMonth = month_before_event(eventMonth, orderDay)[1]
            else:
                orderAndPaymentMonth = eventMonth
            paymentDay = orderDay + 1

            def check_if_needed_zero(x, y):
                if x<10 and y<10:
                    x ="0%s" % str(x)
                    y ="0%s" % str(y)
                elif x<10 and y>=10:
                    x ="0%s" % str(x)
                    y ="%s" % str(y)
                elif x>=10 and y<10:
                    x ="%s" % str(x)
                    y ="0%s" % str(y)
                else:
                    x = str(x)
                    y = str(y)
                return x, y
            tupleWithEventInfo = check_if_needed_zero(eventDay, eventMonth)
            eventDay = tupleWithEventInfo[0]
            eventMonth = tupleWithEventInfo[1]
            tupleWithPaymentAndOrderInfo = check_if_needed_zero(orderDay, paymentDay)
            orderDay = tupleWithPaymentAndOrderInfo[0]
            paymentDay = tupleWithPaymentAndOrderInfo[1]
            orderAndPaymentMonth = check_if_needed_zero(orderAndPaymentMonth, 1)[0]
            global eventTitle
            eventTitle = "%s.%s Event O.%s.%s P.%s.%s\n" % (eventDay, eventMonth, orderDay, orderAndPaymentMonth, paymentDay, orderAndPaymentMonth) 
            # Data to duplicate event to name of sheet
            createEvent.eventDay = eventDay
            createEvent.eventMonth = eventMonth
            createEvent.eventYear = eventYear
            createEvent.orderDay = orderDay
            createEvent.paymentDay = paymentDay
            createEvent.orderAndPaymentMonth = orderAndPaymentMonth
        event_order_payment(eventDayCheck, eventDay, eventMonth)

# Print event description for example so as to create facebook event
    def printEventDescription():
        print('Enter sheet link:')
        sheetLink = str(input())
        print('\nEvent description:')
        print(eventTitle)
        print(f"ORDER: untill {orderDay}.{orderAndPaymentMonth}\n")
        print(f"PAYMENT: untill {paymentDay}.{orderAndPaymentMonth}\n")
        if eventGroup == "Group1":
            print(f"Account: \nXYZ \n(account number)\n")
        else:
            print(f"Account: \nZYX \n(account number)\n")
        print(f'SHEET:\n%s\n' % sheetLink)
        print('event description...\n')
       
    # Loop to choose opeation
    while True:
        print('\nChoose:\n1.Create event\n2.Color rows\n3.Read orders and send them in email')
        toDo = input()
        if toDo == '1':
            chooseGroup()
            createEvent()
            eventDay = createEvent.eventDay
            eventMonth = createEvent.eventMonth
            eventYear = createEvent.eventYear
            orderDay = createEvent.orderDay
            paymentDay = createEvent.paymentDay
            orderAndPaymentMonth = createEvent.orderAndPaymentMonth
            duplicateSheet()
            printEventDescription()
            clearValuesOfSpreadsheet()
        if toDo == '2':
            chooseGroup()
            print('Enter ID of sheet to color:')
            newSheetId= str(input())
            while True:
                print('Enter number of buyer:')
                columnToColor = int(input())
                columnToColor = columnToColor + 6
                colorColumn(columnToColor, newSheetId)
                print('Exit? y/n')
                exitFromColorMethod = str(input())
                if exitFromColorMethod == 'y':
                    break
        if toDo == '3':
            chooseGroup()
            getDataFromSheet()
            print('Create .docx file with order info? y/n')
            createOrNotDocx = str(input())
            if createOrNotDocx == 'y':
                createDocx()
            print('Send order info for e-mail? y/n')
            sendMailOrNot = str(input())
            if sendMailOrNot == 'y':
                sendMail()
                print('Order info was sent.')
            break

if __name__ == '__main__':
    main()
